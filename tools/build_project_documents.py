from __future__ import annotations

from pathlib import Path
from typing import Iterable, Sequence

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "deliverables"
QA = ROOT / "artifacts" / "document_qa"
SRS_PATH = OUT / "DailyCart_Software_Requirements_Specification.docx"
RESEARCH_PATH = OUT / "DailyCart_Research_and_Project_Analysis.docx"

BLUE, DARK_BLUE, NAVY, MUTED = "2E74B5", "1F4D78", "0B2545", "5B6573"
LIGHT_GRAY, NARRATIVE_FILL = "F2F4F7", "F4F6F9"
BLACK, GREEN, AMBER, RED = "000000", "2E7D32", "B7791F", "9B1C1C"


def rgb(value: str) -> RGBColor:
    return RGBColor.from_string(value)


def set_run_font(run, name="Calibri", size=None, color=BLACK, bold=None, italic=None):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = rgb(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_table_borders(table, color="C7CDD6", size=6):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), str(size))
        node.set(qn("w:color"), color)


def set_table_geometry(table, widths_dxa: Sequence[int], indent_dxa=120):
    if sum(widths_dxa) != 9360:
        raise ValueError(f"Table widths must total 9360 DXA, got {sum(widths_dxa)}")
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), "9360")
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = Inches(widths_dxa[idx] / 1440)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths_dxa[idx]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_page_field(paragraph):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar"); begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText"); instr.set(qn("xml:space"), "preserve"); instr.text = " PAGE "
    sep = OxmlElement("w:fldChar"); sep.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t"); text.text = "1"
    end = OxmlElement("w:fldChar"); end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, sep, text, end])
    set_run_font(run, size=9, color=MUTED)


def configure_section(section, running_label: str):
    section.page_width, section.page_height = Inches(8.5), Inches(11)
    section.top_margin = section.bottom_margin = Inches(1)
    section.left_margin = section.right_margin = Inches(1)
    section.header_distance = section.footer_distance = Inches(0.492)
    section.different_first_page_header_footer = True
    hp = section.header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    hp.paragraph_format.space_after = Pt(0)
    set_run_font(hp.add_run(running_label), size=8.5, color=MUTED, bold=True)
    fp = section.footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    fp.paragraph_format.space_before = Pt(0)
    set_run_font(fp.add_run("DailyCart | Page "), size=9, color=MUTED)
    add_page_field(fp)
    section.first_page_header.paragraphs[0].text = ""
    section.first_page_footer.paragraphs[0].text = ""


def setup_styles(doc: Document, preset="standard"):
    if preset == "narrative":
        body_after, body_line = 8, 1.333
        heading_tokens = (("Heading 1", 16, BLUE, 18, 10), ("Heading 2", 13, BLUE, 12, 6), ("Heading 3", 12, DARK_BLUE, 8, 4))
    else:
        body_after, body_line = 6, 1.10
        heading_tokens = (("Heading 1", 16, BLUE, 16, 8), ("Heading 2", 13, BLUE, 12, 6), ("Heading 3", 12, DARK_BLUE, 8, 4))
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size, normal.font.color.rgb = Pt(11), rgb(BLACK)
    normal.paragraph_format.space_before, normal.paragraph_format.space_after = Pt(0), Pt(body_after)
    normal.paragraph_format.line_spacing = body_line
    if preset == "narrative":
        normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    for name, size, color, before, after in heading_tokens:
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size, style.font.color.rgb, style.font.bold = Pt(size), rgb(color), True
        style.paragraph_format.space_before, style.paragraph_format.space_after = Pt(before), Pt(after)
        style.paragraph_format.keep_with_next = style.paragraph_format.keep_together = True
    caption = doc.styles["Caption"]
    caption.font.name = "Calibri"
    caption._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    caption._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    caption.font.size, caption.font.color.rgb, caption.font.italic = Pt(9), rgb(MUTED), True
    caption.paragraph_format.space_before, caption.paragraph_format.space_after = Pt(4), Pt(8)


def add_numbering(doc: Document, kind: str, text_left=720, hanging=360, after=160, line=280) -> int:
    numbering = doc.part.numbering_part.element
    abstract_ids = [int(x.get(qn("w:abstractNumId"))) for x in numbering.findall(qn("w:abstractNum"))]
    num_ids = [int(x.get(qn("w:numId"))) for x in numbering.findall(qn("w:num"))]
    abstract_id, num_id = max(abstract_ids, default=0) + 1, max(num_ids, default=0) + 1
    abstract = OxmlElement("w:abstractNum"); abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType"); multi.set(qn("w:val"), "singleLevel"); abstract.append(multi)
    lvl = OxmlElement("w:lvl"); lvl.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start"); start.set(qn("w:val"), "1"); lvl.append(start)
    fmt = OxmlElement("w:numFmt"); fmt.set(qn("w:val"), "bullet" if kind == "bullet" else "decimal"); lvl.append(fmt)
    lvl_text = OxmlElement("w:lvlText"); lvl_text.set(qn("w:val"), "•" if kind == "bullet" else "%1."); lvl.append(lvl_text)
    suff = OxmlElement("w:suff"); suff.set(qn("w:val"), "tab"); lvl.append(suff)
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs"); tab = OxmlElement("w:tab"); tab.set(qn("w:val"), "num"); tab.set(qn("w:pos"), str(text_left)); tabs.append(tab); p_pr.append(tabs)
    ind = OxmlElement("w:ind"); ind.set(qn("w:left"), str(text_left)); ind.set(qn("w:hanging"), str(hanging)); p_pr.append(ind)
    spacing = OxmlElement("w:spacing"); spacing.set(qn("w:after"), str(after)); spacing.set(qn("w:line"), str(line)); spacing.set(qn("w:lineRule"), "auto"); p_pr.append(spacing)
    lvl.append(p_pr); abstract.append(lvl); numbering.append(abstract)
    num = OxmlElement("w:num"); num.set(qn("w:numId"), str(num_id))
    ref = OxmlElement("w:abstractNumId"); ref.set(qn("w:val"), str(abstract_id)); num.append(ref); numbering.append(num)
    return num_id


def add_list_item(doc, text: str, num_id: int):
    p = doc.add_paragraph()
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl"); ilvl.set(qn("w:val"), "0")
    n = OxmlElement("w:numId"); n.set(qn("w:val"), str(num_id))
    num_pr.extend([ilvl, n]); p._p.get_or_add_pPr().append(num_pr)
    p.add_run(text)
    return p


def add_para(doc, text: str, *, bold=False, italic=False, color=BLACK, size=11, align=None, before=0, after=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before)
    if after is not None: p.paragraph_format.space_after = Pt(after)
    if align is not None: p.alignment = align
    set_run_font(p.add_run(text), size=size, color=color, bold=bold, italic=italic)
    return p


def add_labeled_para(doc, label: str, value: str, after=3):
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(after)
    set_run_font(p.add_run(label + ": "), bold=True, color=NAVY); p.add_run(value)
    return p


def add_callout(doc, label: str, text: str, fill=NARRATIVE_FILL, accent=BLUE):
    table = doc.add_table(rows=1, cols=1)
    set_repeat_table_header(table.rows[0])
    set_table_geometry(table, [9360]); set_table_borders(table, color=fill, size=1)
    cell = table.cell(0, 0); set_cell_shading(cell, fill)
    p = cell.paragraphs[0]; p.paragraph_format.space_before = p.paragraph_format.space_after = Pt(2)
    set_run_font(p.add_run(label + ": "), bold=True, color=accent); p.add_run(text)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_table(doc, headers: Sequence[str], rows: Iterable[Sequence[str]], widths: Sequence[int], *, header_fill=LIGHT_GRAY, font_size=9.5):
    table = doc.add_table(rows=1, cols=len(headers)); set_repeat_table_header(table.rows[0])
    for idx, value in enumerate(headers):
        cell = table.rows[0].cells[idx]; set_cell_shading(cell, header_fill)
        p = cell.paragraphs[0]; p.paragraph_format.space_before = p.paragraph_format.space_after = Pt(0)
        set_run_font(p.add_run(str(value)), size=font_size, color=NAVY, bold=True)
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            p = cells[idx].paragraphs[0]; p.paragraph_format.space_before = p.paragraph_format.space_after = Pt(0); p.paragraph_format.line_spacing = 1.05
            set_run_font(p.add_run(str(value)), size=font_size)
    set_table_geometry(table, widths); set_table_borders(table)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


def add_heading(doc, text: str, level=1):
    p = doc.add_heading(text, level=level); p.paragraph_format.keep_with_next = True; return p


def add_cover_srs(doc):
    add_para(doc, "SOFTWARE REQUIREMENTS SPECIFICATION", bold=True, size=10.5, color=BLUE, after=12)
    add_para(doc, "DailyCart", bold=True, size=30, color=NAVY, after=5)
    add_para(doc, "Multi-vendor Grocery Delivery Platform", size=15, color=MUTED, after=24)
    add_callout(doc, "Document purpose", "A complete, implementation-aligned specification for the DailyCart web platform, REST API, and Flutter mobile client.")
    doc.add_paragraph().paragraph_format.space_after = Pt(8)
    for label, value in [("Version", "1.0"), ("Status", "Baseline specification derived from the repository snapshot"), ("Prepared", "18 July 2026"), ("Currency and locale", "Sri Lankan Rupees (LKR); English-first interface"), ("Primary platforms", "Laravel 12 web application and API; Flutter mobile application"), ("Verification snapshot", "54 automated tests passed (241 assertions); Vite production build passed")]:
        add_labeled_para(doc, label, value)
    add_para(doc, "DailyCart Project Documentation", size=9.5, color=MUTED, italic=True, before=34, after=0)
    doc.add_page_break()


def add_cover_research(doc):
    add_para(doc, "TECHNICAL RESEARCH AND PROJECT ANALYSIS", bold=True, size=10, color=AMBER, align=WD_ALIGN_PARAGRAPH.CENTER, before=70, after=18)
    add_para(doc, "DailyCart", bold=True, size=30, color=NAVY, align=WD_ALIGN_PARAGRAPH.CENTER, after=7)
    add_para(doc, "Architecture, Feasibility, Readiness, and Improvement Roadmap", size=15, color=DARK_BLUE, align=WD_ALIGN_PARAGRAPH.CENTER, after=28)
    add_para(doc, "A repository-based evaluation of a Sri Lankan multi-vendor grocery delivery platform", size=11, color=MUTED, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, after=80)
    add_para(doc, "Prepared from the project snapshot and authoritative technical, market, payment, security, and regulatory sources", size=10, color=MUTED, align=WD_ALIGN_PARAGRAPH.CENTER, after=8)
    add_para(doc, "18 July 2026", size=11, color=NAVY, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, after=0)
    doc.add_page_break()


def font(size, bold=False):
    for path in [Path("C:/Windows/Fonts/arialbd.ttf" if bold else "C:/Windows/Fonts/arial.ttf"), Path("C:/Windows/Fonts/calibrib.ttf" if bold else "C:/Windows/Fonts/calibri.ttf")]:
        if path.exists(): return ImageFont.truetype(str(path), size)
    return ImageFont.load_default()


def rounded_box(draw, xy, fill, outline, title, lines):
    draw.rounded_rectangle(xy, radius=22, fill=fill, outline=outline, width=3)
    x1, y1, _, _ = xy; draw.text((x1 + 24, y1 + 20), title, font=font(34, True), fill="#0B2545")
    y = y1 + 72
    for line in lines: draw.text((x1 + 26, y), line, font=font(25), fill="#25313D"); y += 36


def arrow(draw, start, end, color="#66788A"):
    draw.line([start, end], fill=color, width=6); x, y = end; draw.polygon([(x, y), (x - 18, y - 10), (x - 18, y + 10)], fill=color)


def create_architecture_figure(path: Path):
    img = Image.new("RGB", (1800, 1000), "#FFFFFF"); d = ImageDraw.Draw(img)
    d.text((60, 35), "DailyCart current-state architecture", font=font(44, True), fill="#0B2545")
    rounded_box(d, (70, 145, 440, 390), "#F2F7FC", "#8EB6D9", "Clients", ["Blade + Tailwind web", "Flutter mobile", "Admin / vendor / rider"])
    rounded_box(d, (605, 110, 1195, 425), "#F8FAFC", "#7D95AA", "Laravel 12 application", ["Routes + middleware + policies", "Controllers + Form Requests", "35 domain services", "Queues, scheduler, notifications"])
    rounded_box(d, (1360, 145, 1730, 390), "#F3F8F3", "#8BB58B", "Data", ["MySQL 8", "Eloquent models", "Filesystem / S3", "Encrypted backups"])
    rounded_box(d, (335, 600, 780, 880), "#FFF9ED", "#D8B15B", "External services", ["PayHere payments", "Google Maps", "SMTP / SMS / push", "Firebase mobile messaging"])
    rounded_box(d, (1020, 600, 1465, 880), "#F7F4FA", "#A992BE", "Operations", ["Database queue worker", "Daily scheduler", "CI: tests, build, audits", "Health endpoint /up"])
    arrow(d, (440, 265), (605, 265)); arrow(d, (1195, 265), (1360, 265)); arrow(d, (780, 735), (920, 510)); arrow(d, (1020, 735), (920, 510))
    d.text((690, 470), "HTTPS / session / Sanctum tokens", font=font(25, True), fill="#526579")
    d.text((60, 940), "Evidence: repository routes, controllers, services, configuration, migrations, mobile services, and CI workflow.", font=font(22), fill="#596A78")
    img.save(path)


def create_readiness_figure(path: Path):
    scores = [("Web feature breadth", 4.5, GREEN), ("Commerce integrity", 4.4, GREEN), ("Security controls", 3.8, BLUE), ("Automated verification", 3.7, BLUE), ("Operational readiness", 3.4, AMBER), ("Mobile API alignment", 1.8, RED)]
    img = Image.new("RGB", (1700, 860), "#FFFFFF"); d = ImageDraw.Draw(img)
    d.text((55, 35), "Repository-based readiness assessment (0-5)", font=font(42, True), fill="#0B2545"); y = 135
    for label, score, color in scores:
        d.text((60, y), label, font=font(29, True), fill="#25313D"); d.rounded_rectangle((520, y + 4, 1520, y + 42), radius=16, fill="#E9EDF2")
        width = int(1000 * score / 5); d.rounded_rectangle((520, y + 4, 520 + width, y + 42), radius=16, fill="#" + color); d.text((1540, y), f"{score:.1f}", font=font(29, True), fill="#25313D"); y += 105
    d.text((60, 790), "Scores synthesize implemented scope, tests, configuration, and observed contract gaps; they are analytical, not certification results.", font=font(22), fill="#596A78")
    img.save(path)


def add_figure(doc, path: Path, caption: str, width=6.25):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_before = Pt(5); p.paragraph_format.space_after = Pt(2)
    inline_shape = p.add_run().add_picture(str(path), width=Inches(width))
    inline_shape._inline.docPr.set("descr", caption)
    inline_shape._inline.docPr.set("title", caption.split(".", 1)[0])
    cp = doc.add_paragraph(caption, style="Caption"); cp.alignment = WD_ALIGN_PARAGRAPH.CENTER


def req_rows(prefix: str, items: Sequence[tuple[str, str]]) -> list[tuple[str, str, str]]:
    return [(f"{prefix}-{i:02d}", text, priority) for i, (text, priority) in enumerate(items, 1)]


def add_requirements(doc, title: str, prefix: str, items: Sequence[tuple[str, str]]):
    add_heading(doc, title, 2); add_table(doc, ["ID", "Requirement", "Priority"], req_rows(prefix, items), [1150, 7110, 1100], font_size=9.2)


def build_srs(architecture_path: Path):
    doc = Document(); setup_styles(doc); configure_section(doc.sections[0], "DailyCart Software Requirements Specification")
    doc.core_properties.title, doc.core_properties.subject, doc.core_properties.author = "DailyCart Software Requirements Specification", "Complete implementation-aligned SRS", "DailyCart Project Team"
    bullet_id, number_id = add_numbering(doc, "bullet"), add_numbering(doc, "decimal")
    add_cover_srs(doc)
    add_heading(doc, "Document Control", 1)
    add_table(doc, ["Field", "Value"], [("Document identifier", "DC-SRS-001"), ("Version", "1.0"), ("Status", "Repository-derived baseline"), ("Prepared date", "18 July 2026"), ("System release observed", "Laravel Framework 12.64.0; Flutter app version 1.0.0+1"), ("Evidence window", "Source code, configuration, migrations, routes, project guides, tests, and build output available in the workspace"), ("Validation result", "54 PHPUnit tests passed with 241 assertions; Vite production build succeeded")], [2200, 7160])
    add_heading(doc, "Contents", 1)
    for entry in ["1. Introduction and conventions", "2. Product overview", "3. Stakeholders and actors", "4. System context and architecture", "5. Representative use cases", "6. Functional requirements", "7. Business rules", "8. Data requirements", "9. External interface requirements", "10. Non-functional requirements", "11. Security, privacy, and compliance", "12. Deployment and operations", "13. Acceptance and verification", "14. Dependencies, risks, and open decisions", "Appendices: states, API baseline, traceability, glossary"]: add_list_item(doc, entry, number_id)
    add_heading(doc, "Executive Summary", 1)
    add_para(doc, "DailyCart is a Sri Lanka-focused, multi-vendor grocery commerce and delivery platform. The observed implementation combines a Laravel 12 monolith, server-rendered Blade/Tailwind web interfaces, a versioned Sanctum-protected REST API, MySQL persistence, queued notification channels, scheduled subscription processing, and a Flutter client. It supports five business roles: Super Admin, Admin, Vendor, Rider, and Customer.")
    add_para(doc, "This specification records the intended product behavior and the requirements necessary to operate the platform safely. Requirements marked Must define the baseline release; Should items are expected for production quality; Could items are controlled enhancements. Where the Flutter client currently calls endpoints not exposed by the API, the requirement is retained as a target and the mismatch is recorded as an implementation gap rather than represented as completed functionality.")
    add_callout(doc, "Baseline finding", "The web platform is feature-rich and verified by automated tests. The highest-priority delivery risk is incomplete alignment between the Flutter service contract and the current /api/v1 route surface.")
    add_heading(doc, "1. Introduction", 1); add_heading(doc, "1.1 Purpose", 2)
    add_para(doc, "This document defines functional, data, interface, quality, security, operational, and acceptance requirements for DailyCart. It is intended to guide product owners, developers, testers, database administrators, deployment engineers, academic reviewers, and future maintainers.")
    add_heading(doc, "1.2 Scope", 2)
    for item in ["Public discovery pages, account registration, authentication, verification, and role-directed dashboards.", "Multi-vendor catalog, inventory, promotions, coupons, advertisements, cart, wishlist, checkout, and order splitting.", "Payments in LKR using cash on delivery, PayHere card flow, bank transfer placeholder, and customer wallet.", "Delivery scheduling, rider assignment, location updates, proof of delivery, earnings, and reporting.", "Refunds, loyalty, recurring subscriptions, reviews, support tickets, notifications, content, and platform administration.", "A REST API and Flutter application for customer, vendor, and rider experiences."]: add_list_item(doc, item, bullet_id)
    add_heading(doc, "1.3 Product Boundaries", 2); add_para(doc, "DailyCart is an application platform and marketplace workflow system. It does not itself operate banking networks, card authorization infrastructure, mapping datasets, mobile app stores, telecom gateways, or third-party mail delivery. Those capabilities are external dependencies governed through adapters and configuration.")
    add_heading(doc, "1.4 Requirement Language", 2)
    add_table(doc, ["Term", "Meaning"], [("Must", "Mandatory for the stated baseline and acceptance."), ("Should", "Expected for production readiness unless an approved exception exists."), ("Could", "Desirable enhancement that may be scheduled after the baseline."), ("System", "The complete DailyCart solution, including web, API, worker, scheduler, database, and mobile client.")], [1800, 7560])
    add_heading(doc, "2. Product Overview", 1); add_heading(doc, "2.1 Product Perspective", 2)
    add_para(doc, "DailyCart is implemented as a modular monolith. HTTP routes and middleware establish the boundary; controllers coordinate requests; Form Requests validate and authorize input; policies and role middleware enforce ownership and permissions; domain services handle transactional commerce rules; Eloquent models persist state; background jobs and the scheduler handle deferred work.")
    add_heading(doc, "2.2 Business Objectives", 2)
    for item in ["Enable customers to find and purchase groceries from multiple approved vendors in a single cart.", "Give vendors controlled catalog, inventory, order, promotion, refund, earning, and reporting capabilities.", "Coordinate approved riders through an auditable delivery lifecycle with proof of completion.", "Provide administrators with governance, finance, moderation, analytics, reporting, and maintenance controls.", "Support local payment and delivery expectations through LKR pricing, PayHere, wallet, cash on delivery, and Sri Lankan location structures."]: add_list_item(doc, item, bullet_id)
    add_heading(doc, "2.3 Operating Environment", 2)
    add_table(doc, ["Layer", "Baseline"], [("Server", "PHP 8.2 or newer; Laravel 12; Composer 2; Linux or equivalent production host"), ("Data", "MySQL 8; database-backed sessions, cache, and queue supported"), ("Web UI", "Modern browser with JavaScript; Blade, Tailwind CSS, Alpine.js, Axios, Vite assets"), ("Mobile", "Flutter stable, Dart 3.10 or newer; Android and iOS targets"), ("Storage", "Private/local or S3-compatible backup disk; public product and proof media storage"), ("Operations", "Queue worker, scheduler/cron, HTTPS reverse proxy, environment secret management")], [1900, 7460])
    add_heading(doc, "2.4 Constraints and Assumptions", 2)
    for item in ["Prices, fees, wallets, refunds, and reporting use LKR and decimal monetary storage.", "The current interface is English-first; multilingual localization is not implemented as a release baseline.", "Vendors and riders require administrative approval before protected operations.", "PayHere callbacks require a publicly reachable HTTPS endpoint and valid merchant credentials.", "Google Maps, SMTP, SMS, push, Firebase, object storage, queues, and app-store accounts are externally provisioned.", "A customer cart may contain items from multiple vendors, but persisted orders are split so each order belongs to one vendor."]: add_list_item(doc, item, bullet_id)
    add_heading(doc, "3. Stakeholders and Actors", 1)
    add_table(doc, ["Actor", "Primary goals", "Access boundary"], [("Visitor", "Browse public products, categories, offers, policies, contact, and newsletter; register or log in", "Public data only"), ("Customer", "Manage profile and addresses; shop, pay, track, review, request support/refunds, manage loyalty and subscriptions", "Own customer resources"), ("Vendor", "Manage store products, inventory, variants, orders, promotions, coupons, reviews, earnings, and reports", "Approved vendor and owned resources"), ("Rider", "Review assignments, update delivery state/location, submit proof, and view earnings/reports", "Verified rider and assigned deliveries"), ("Admin", "Moderate marketplace operations, finance, deliveries, content, support, promotions, reports, and approvals", "Administrative functions; excludes protected Super Admin ownership"), ("Super Admin", "Manage admins, settings, roles, logs, backups, maintenance, and platform-wide controls", "Highest privileged role"), ("External provider", "Supply payment, maps, email, SMS, push, cloud storage, and app distribution", "Contracted integration boundary")], [1450, 5000, 2910], font_size=9.2)
    add_heading(doc, "4. System Context and Architecture", 1); add_figure(doc, architecture_path, "Figure 1. DailyCart current-state component and integration context.")
    add_heading(doc, "4.1 Architectural Responsibilities", 2)
    add_table(doc, ["Component", "Responsibility"], [("Web routes and controllers", "Serve public, authenticated, customer, vendor, rider, admin, super-admin, and integration workflows."), ("REST API v1", "Expose token-authenticated mobile endpoints with scoped abilities and verification gates."), ("Domain services", "Apply pricing, stock, transaction, delivery, payment, refund, loyalty, notification, reporting, and backup rules."), ("MySQL", "Persist identities, RBAC, profiles, catalog, commerce, delivery, finance, engagement, integrations, and operational queues."), ("Queue and scheduler", "Defer communication and recurring order work; prune expired API tokens."), ("Flutter client", "Provide cross-platform customer, vendor, and rider experiences through service/provider/screen layers.")], [2500, 6860])
    add_heading(doc, "5. Representative Use Cases", 1)
    use_cases = [("UC-01", "Customer registration creates User, Customer, role assignment, verification state, and location."), ("UC-02", "Vendor onboarding remains pending until Admin approval or rejection."), ("UC-03", "Multi-vendor checkout quotes, discounts, schedules, pays, and creates one order per vendor transactionally."), ("UC-04", "Vendor confirms/packs; Admin assigns verified rider; Rider completes delivery with proof."), ("UC-05", "PayHere card flow validates a signed server callback before payment state changes."), ("UC-06", "Wallet payment debits atomically; approved eligible refund credits wallet."), ("UC-07", "Scheduler queues due subscriptions and generates pending recurring orders."), ("UC-08", "Customer reviews delivered purchases and exchanges owned support ticket replies."), ("UC-09", "Super Admin manages admins, roles, settings, logs, backups, and caches."), ("UC-10", "Mobile role workflow uses Sanctum and only documented API endpoints.")]
    add_table(doc, ["ID", "Use case summary"], use_cases, [1200, 8160], font_size=9.4)
    add_heading(doc, "6. Functional Requirements", 1)
    add_requirements(doc, "6.1 Identity, Authentication, and Access Control", "FR-AUTH", [("The system must support customer, vendor, and rider self-registration with role-specific profile creation.", "Must"), ("The system must support email/password login, logout, password reset, password confirmation, and password change.", "Must"), ("The system must support email verification and phone OTP verification before sensitive API workflows.", "Must"), ("The system must optionally challenge web login with OTP when configured.", "Should"), ("The system must assign one primary DailyCart role and synchronize compatible Spatie role membership.", "Must"), ("The system must redirect authenticated users to the dashboard for their primary role.", "Must"), ("The system must deny cross-role route access and enforce resource ownership through policies.", "Must"), ("The system must hold new vendors and riders in pending state until an Admin or Super Admin decision.", "Must"), ("The system must issue expiring Sanctum tokens with abilities appropriate to the authenticated role.", "Must"), ("The system must rate limit API login, registration, and OTP operations.", "Must"), ("The system should log security-relevant account and administrative events.", "Should"), ("The system must allow users to update profiles and delete their own account subject to password confirmation.", "Must")])
    add_requirements(doc, "6.2 Catalog, Vendor, and Inventory", "FR-CAT", [("Admins must create, update, deactivate, and hierarchically organize categories and brands.", "Must"), ("Approved vendors must create, view, update, and delete only their own products.", "Must"), ("Products must support category, brand, SKU, description, unit, price, discount price, subscription eligibility, and status.", "Must"), ("Products must support a primary image, gallery images, and alt text-compatible metadata.", "Must"), ("Products must support active variants and variant-level inventory where configured.", "Must"), ("Vendor product changes must return the product to an approval workflow when policy requires.", "Must"), ("Admins must approve, reject, feature, activate, or deactivate products.", "Must"), ("Customer catalog queries must expose only approved products from active categories and approved vendors.", "Must"), ("The system must prevent ordering quantities greater than the lesser available product and variant stock.", "Must"), ("The system should notify vendors when stock reaches the configurable low-stock threshold.", "Should"), ("Catalog browsing should support search, category, price, rating, availability, brand, and sorting filters consistently across web and API.", "Should")])
    add_requirements(doc, "6.3 Cart, Wishlist, Pricing, and Checkout", "FR-CHK", [("The system must maintain one active cart per customer and prevent access to another customer's cart items.", "Must"), ("Customers must add, update, remove, and clear eligible products and variants.", "Must"), ("The system must refresh promotional pricing and revalidate stock before presenting checkout totals.", "Must"), ("Customers must add unique wishlist items and move an item from wishlist to cart.", "Must"), ("Checkout must calculate subtotal, coupon discount, loyalty discount, delivery fee, service charge, tax, and grand total in LKR.", "Must"), ("The system must apply the best active applicable promotion without producing a negative item price.", "Must"), ("Coupon validation must enforce status, validity dates, minimum amount, vendor scope, global limit, customer limit, discount type, and maximum discount.", "Must"), ("A coupon may be applied to at most one vendor order group during a single multi-vendor checkout.", "Must"), ("Loyalty redemption must not exceed the customer's balance or make any order total negative.", "Must"), ("The customer must select a delivery time at least 30 minutes after placement.", "Must"), ("Checkout must split cart items into one order per vendor while preserving a single user-facing transaction outcome.", "Must"), ("Checkout must lock stock rows, atomically decrement stock, create order items/payment/delivery records, and convert the cart.", "Must"), ("A quote and the created order must use the same pricing rules and delivery configuration.", "Must")])
    add_requirements(doc, "6.4 Order and Delivery Lifecycle", "FR-ORD", [("Customers must list, view, track, and obtain a receipt only for their own orders.", "Must"), ("Customers must cancel an order only while its state is pending, subject to the authoritative cancellation policy.", "Must"), ("Vendors must view and manage only orders belonging to their vendor profile.", "Must"), ("A vendor must confirm a pending order and mark only a confirmed order as packed.", "Must"), ("Admins must filter all orders, update permitted statuses, and assign a verified rider only after packing.", "Must"), ("The system must update the assigned rider to delivering and notify the customer and rider.", "Must"), ("A rider must access only assigned deliveries and may record acceptance.", "Must"), ("A rider must progress a delivery from assigned to picked up to on the way to delivered.", "Must"), ("Delivery completion must store timestamp, required proof image, optional signature, and note.", "Must"), ("Delivery completion must synchronize order state, COD payment state, rider availability, loyalty earning, and notifications atomically.", "Must"), ("A rider may mark an assigned, picked-up, or on-the-way delivery failed with a reason.", "Must"), ("The system should retain status history sufficient for customer tracking and audit.", "Should"), ("Delivery fees must be configurable by district/distance/quantity logic and used consistently by quote, order, and payment.", "Must")])
    add_requirements(doc, "6.5 Payments, Wallets, Refunds, and Finance", "FR-PAY", [("Each order must have at most one payment record and an allowed method: cash on delivery, card, bank transfer, or wallet.", "Must"), ("Payment records must mirror order subtotal, discounts, fees, grand total, amount, currency, status, and references.", "Must"), ("Wallet checkout must atomically lock the customer balance, prevent overdraft, debit once, create a ledger entry, and mark payment paid.", "Must"), ("Cash on delivery must remain pending until successful delivery completion.", "Must"), ("The PayHere flow must generate a correct checkout request and update state only after validating the signed server callback.", "Must"), ("The system must not treat the PayHere return redirect as authoritative payment confirmation.", "Must"), ("Bank transfer and test simulation behavior must be disabled or explicitly controlled in production.", "Should"), ("Customer wallet top-up must use a real authorized payment source in production; placeholder crediting must not be publicly enabled.", "Must"), ("Customers may request a refund only for delivered, paid orders and for an amount within the unrefunded paid balance.", "Must"), ("Admin approval must credit the wallet, mark payment refunded, mark order refunded, reverse applicable loyalty, and notify the customer transactionally.", "Must"), ("Finance reports must expose revenue, fees, payouts, refunds, COD pending amounts, and paid-order counts with date filtering.", "Must"), ("Vendor earnings must exclude refunded orders and apply commission; rider earnings must include delivered work only.", "Must")])
    add_requirements(doc, "6.6 Marketing, Loyalty, Reviews, Support, and Content", "FR-ENG", [("Admins and authorized vendors must manage scoped coupons and promotions.", "Must"), ("Admins must manage advertisements and ensure only active, in-window, approved content appears on the storefront.", "Must"), ("The system must record coupon redemptions and basic promotion/advertisement performance counters.", "Should"), ("Admins must configure loyalty earn and redemption values.", "Must"), ("The system must earn loyalty points once for a delivered order and record redemption/reversal ledger entries.", "Must"), ("Customers must review only a product purchased in their own delivered order and only once per order/product.", "Must"), ("Review ratings must be 1 through 5; Admins must hide or delete inappropriate content; Vendors may view only reviews of their products.", "Must"), ("Any authenticated user must create and access only their own support tickets, with optional owned-order association.", "Must"), ("Admins must assign, prioritize, change state, and reply to all support tickets.", "Must"), ("Users must view their notifications and mark them read or unread; Admins may view platform notifications.", "Must"), ("The system should deliver configured email, SMS, WhatsApp, and push channels asynchronously after transaction commit.", "Should"), ("Admins must manage public policy/content pages, contact messages, and newsletter subscriptions.", "Must")])
    add_requirements(doc, "6.7 Subscriptions and Scheduled Orders", "FR-SUB", [("Customers must create subscriptions only for approved, subscription-eligible products and active variants with sufficient stock.", "Must"), ("Subscriptions must capture frequency, quantity, unit price, total, address, preferred time, dates, payment method, notes, and next delivery date.", "Must"), ("Customers must view, update, pause, resume, and cancel their own subscriptions subject to state rules.", "Must"), ("The scheduler must identify due active subscriptions daily and dispatch one idempotent generation job per due subscription.", "Must"), ("Recurring generation must create a vendor-specific pending order, item, payment, and delivery record and advance the next date transactionally.", "Must"), ("Stock for a subscription-generated order must be validated and reduced when the vendor confirms the order.", "Must"), ("Generation failures must preserve the subscription, record the reason, and notify customer and vendor.", "Must"), ("Customers, vendors, and admins must view scheduled orders within their authorization scope.", "Must"), ("A customer may cancel only an owned pending scheduled order.", "Must")])
    add_requirements(doc, "6.8 Administration, Reporting, and Maintenance", "FR-ADM", [("Admins must access operational dashboards, analytics, customer, vendor, rider, product, order, delivery, refund, finance, support, marketing, and report functions.", "Must"), ("Super Admins must create, update, suspend, and remove Admin accounts without permitting ordinary Admins to modify Super Admin ownership.", "Must"), ("Super Admins must manage role permissions and platform settings.", "Must"), ("Super Admins must review activity, API integration, and security logs.", "Must"), ("Authorized users must export reports in controlled, auditable formats.", "Should"), ("Super Admins must trigger encrypted database backups and securely download authorized backup files.", "Must"), ("Backup retention must support both count and age policies on a private durable disk.", "Must"), ("The platform must provide safe compiled-cache maintenance without deleting application or user data.", "Must"), ("Administrative actions with financial, identity, authorization, or moderation impact should be audit logged.", "Should"), ("The system must expose a health endpoint suitable for deployment checks.", "Must")])
    add_requirements(doc, "6.9 REST API and Mobile Contract", "FR-API", [("The public API must be versioned under /api/v1 and return consistent JSON envelopes and validation errors.", "Must"), ("The API must expose public category and approved-product browsing.", "Must"), ("Authenticated API routes must require Sanctum and abilities appropriate to authentication, profile, verification, customer, vendor, or rider use.", "Must"), ("The API must enforce email and phone verification for protected commerce and operational actions.", "Must"), ("Customer API must expose cart, quote, order creation, order list, and order detail as implemented.", "Must"), ("Vendor API must expose an overview, owned orders, and vendor wallet at minimum.", "Must"), ("Rider API must expose assigned deliveries, delivery detail, permitted status updates, and location reporting.", "Must"), ("The Flutter base URL must include the same /api/v1 prefix or the client must prepend v1 centrally.", "Must"), ("Every mobile service path must have a matching documented backend route, HTTP method, request schema, response schema, and authorization rule before release.", "Must"), ("Unimplemented mobile modules must be feature-gated or clearly excluded from release rather than failing at runtime.", "Must"), ("The API should publish an OpenAPI specification and contract tests covering all mobile endpoints.", "Should"), ("Mobile tokens must be stored only in secure platform storage and removed on logout or authentication failure.", "Must")])
    add_heading(doc, "7. Business Rules", 1)
    rules = [("BR-01", "All monetary values are denominated and displayed in LKR; persisted money uses fixed precision."), ("BR-02", "A product is orderable only when product/category/vendor approval and stock conditions pass."), ("BR-03", "A variant must belong to the selected product and be active."), ("BR-04", "Product and variant stock must never become negative; stock changes occur under row locks."), ("BR-05", "A cart is customer-owned and active until converted or abandoned."), ("BR-06", "A multi-vendor cart is split into vendor-specific orders."), ("BR-07", "Scheduled delivery is no earlier than placement plus 30 minutes."), ("BR-08", "Service charge defaults to 2 percent but is platform-configurable."), ("BR-09", "Delivery charges are configurable and may use district, distance, quantity, and customer rules."), ("BR-10", "Coupon discount cannot exceed eligible subtotal plus free-delivery value and may be capped."), ("BR-11", "Loyalty redemption cannot exceed available points or reduce an order below zero."), ("BR-12", "Cancellation follows the authoritative service state and actor policy."), ("BR-13", "Vendor progression is pending -> confirmed -> packed."), ("BR-14", "Rider assignment requires a packed order and a verified rider."), ("BR-15", "Delivery progression is assigned -> picked up -> on the way -> delivered, with failure before delivery."), ("BR-16", "COD payment becomes paid only when delivery is completed."), ("BR-17", "A wallet balance cannot become negative and every balance change has a ledger entry."), ("BR-18", "Refund requests require delivered and paid orders; approved totals cannot exceed paid value."), ("BR-19", "Loyalty is earned once per delivered order and may be reversed on refund or cancellation."), ("BR-20", "A customer may review a purchased product once per delivered order."), ("BR-21", "Vendors and riders are operational only after approval/verification."), ("BR-22", "Recurring order generation advances according to subscription frequency."), ("BR-23", "PayHere success is accepted only after callback signature verification."), ("BR-24", "Backup data is encrypted and stored on a non-public disk with a separately managed key."), ("BR-25", "Data access remains scoped by role, ability, policy, and record ownership.")]
    add_table(doc, ["ID", "Rule"], rules, [1200, 8160], font_size=9.2)
    add_heading(doc, "8. Data Requirements", 1); add_heading(doc, "8.1 Core Data Domains", 2)
    add_table(doc, ["Domain", "Principal entities", "Integrity requirement"], [("Identity and RBAC", "users, roles, permissions, pivots, tokens", "Unique identities; hashed credentials; scoped roles and tokens"), ("Profiles and location", "customers, vendors, riders, admins, addresses, districts, cities, zones", "One role profile per user; validated active location references"), ("Catalog", "categories, brands, products, images, variants, inventory", "Vendor ownership; approval; non-negative stock"), ("Shopping", "carts, items, wishlists, coupons, redemptions", "Customer ownership; uniqueness rules"), ("Orders", "orders, items, status history, schedules", "Unique number; vendor-specific order; price snapshot"), ("Finance", "payments, gateway transactions, wallets, vendor wallets, refunds", "One payment/order; auditable ledger; bounded refunds"), ("Delivery", "deliveries, rider locations, proofs, fees", "One delivery/order; rider ownership; proof"), ("Engagement", "reviews, notifications, tickets/replies, loyalty", "Owned records; moderation; ledger history"), ("Marketing/content", "promotions, ads, newsletter, contacts, settings", "Validity windows and approval state"), ("Operations", "jobs, failed jobs, sessions, cache, activity/API logs", "Retention, privacy, timestamps, failure visibility")], [1700, 4200, 3460], font_size=8.9)
    add_heading(doc, "8.2 Data Quality and Retention", 2)
    for item in ["Primary and foreign keys must preserve referential integrity; delete behavior must match business retention needs.", "Order items must retain product name, quantity, unit price, and line total as an order-time snapshot.", "Timestamps must be stored consistently and rendered in the configured timezone.", "Soft deletion must not remove evidence needed for finance, refunds, audit, or support.", "Retention schedules must cover personal data, logs, location, proofs, messages, backups, and financial records.", "Sensitive exports and backups must be access-controlled, encrypted in transit, and auditable."]: add_list_item(doc, item, bullet_id)
    add_heading(doc, "9. External Interface Requirements", 1)
    add_table(doc, ["Interface", "Requirement"], [("Responsive web UI", "Current desktop/mobile browsers, semantic navigation, validated forms, safe errors, role dashboards."), ("Flutter UI", "Android/iOS layouts with secure authentication, connectivity/error states, accessible controls, and API-matched features."), ("REST/JSON", "HTTPS, /api/v1, bearer tokens, abilities, pagination, validation errors, stable resources."), ("PayHere", "Configured sandbox/live endpoints, LKR, checkout hash, public notify URL, signature verification, idempotency."), ("Google Maps", "Separate restricted browser/mobile/server keys, HTTPS, server proxying where appropriate."), ("Email/SMS/push", "Secrets, post-commit queues, safe retries, bounded diagnostic logging."), ("Storage", "Public approved media; private backups and exports."), ("MySQL", "Transactions, constraints, indexes, fixed-precision money, and locks for competing updates.")], [1750, 7610], font_size=9.2)
    add_heading(doc, "10. Non-Functional Requirements", 1)
    nfr_groups = {
        "10.1 Security": [("NFR-SEC-01", "All production browser, API, callback, and administrative traffic must use HTTPS.", "Must"), ("NFR-SEC-02", "Passwords must use adaptive hashing and never be logged or returned.", "Must"), ("NFR-SEC-03", "Server-side authorization must combine route, ability, policy, and ownership checks.", "Must"), ("NFR-SEC-04", "Login, registration, OTP, and abuse-prone endpoints must be rate limited and observable.", "Must"), ("NFR-SEC-05", "Uploads must be type/size validated, renamed, non-executable, and traversal-safe.", "Must"), ("NFR-SEC-06", "Secrets must come from deployment secret management and be excluded from source/logs.", "Must"), ("NFR-SEC-07", "Payment callbacks must validate authenticity, amount, currency, ownership, and replay/idempotency.", "Must"), ("NFR-SEC-08", "CI dependency audits must remediate or formally accept high-severity findings.", "Should")],
        "10.2 Performance and Capacity": [("NFR-PERF-01", "Under agreed baseline load, 95 percent of cached reads should complete within 1.5 seconds at the server boundary.", "Should"), ("NFR-PERF-02", "Checkout quote and order placement should complete within 3 seconds excluding gateway redirects.", "Should"), ("NFR-PERF-03", "Catalog, order, delivery, notification, and report queries must paginate and use selective indexes.", "Must"), ("NFR-PERF-04", "Slow communication and recurring work must execute through managed workers.", "Must"), ("NFR-PERF-05", "Capacity tests must define concurrent shoppers, checkout contention, rider updates, and report volume.", "Must")],
        "10.3 Reliability and Availability": [("NFR-REL-01", "Transactional workflows must fully complete or roll back without partial commerce state.", "Must"), ("NFR-REL-02", "Queue jobs and gateway callbacks must be retry-safe and idempotent.", "Must"), ("NFR-REL-03", "Production must supervise web, queue, and scheduler processes and alert on repeated failure.", "Must"), ("NFR-REL-04", "Health monitoring should include database, storage, and queue readiness beyond /up.", "Should"), ("NFR-REL-05", "Availability and maintenance windows must be formally agreed before launch.", "Must")],
        "10.4 Usability and Accessibility": [("NFR-UX-01", "Critical workflows must work at 360-pixel width and desktop layouts without loss of function.", "Must"), ("NFR-UX-02", "Forms must provide field errors, safe retained input, recovery, and no sensitive internals.", "Must"), ("NFR-UX-03", "State labels must be consistent across web, API, email, and mobile.", "Must"), ("NFR-UX-04", "Web/mobile should meet WCAG 2.2 AA expectations.", "Should"), ("NFR-UX-05", "The platform should be localization-ready for Sinhala and Tamil.", "Should")],
        "10.5 Maintainability and Portability": [("NFR-MNT-01", "Business rules must remain in services/policies/requests rather than duplicated in clients/controllers.", "Must"), ("NFR-MNT-02", "Schema changes must use reversible migrations compatible with deployment.", "Must"), ("NFR-MNT-03", "CI must execute tests, formatting, asset build, and dependency audits.", "Must"), ("NFR-MNT-04", "The API should have OpenAPI documentation and automated client contract verification.", "Should"), ("NFR-MNT-05", "Environment configuration must be externalized and deployment-validated.", "Must")],
        "10.6 Backup, Recovery, and Observability": [("NFR-OPS-01", "Backups must be streamed, encrypted, private, retention-limited, and restore-tested.", "Must"), ("NFR-OPS-02", "RPO and RTO must be approved and validated by drills.", "Must"), ("NFR-OPS-03", "Logs need correlation without passwords, tokens, secrets, full payment data, or excess personal data.", "Must"), ("NFR-OPS-04", "Metrics should cover errors/latency, queues, callbacks, checkout failures, contention, and backup age.", "Should"), ("NFR-OPS-05", "Deployment must include database-compatible rollback and worker restart instructions.", "Must")],
    }
    for title, rows in nfr_groups.items(): add_heading(doc, title, 2); add_table(doc, ["ID", "Requirement", "Priority"], rows, [1300, 6960, 1100], font_size=9.1)
    add_heading(doc, "11. Security, Privacy, and Compliance", 1)
    add_para(doc, "DailyCart processes identity, contact, address, location, purchase, payment reference, support, review, device-token, and operational-log data. Production operation therefore requires a data inventory, lawful-purpose analysis, prominent notices, purpose limitation, access controls, retention, data-subject request handling, incident response, and processor agreements.")
    add_table(doc, ["Control area", "Minimum control"], [("Personal data", "Inventory fields/purposes; minimize; document retention; support applicable access, correction, erasure, and objection workflows."), ("E-commerce transparency", "Display seller/product identity, total fees, delivery, cancellation/refund rules, privacy notice, and receipt."), ("Payments", "Keep card entry with PayHere; store gateway references/status only; verify callbacks."), ("Location", "Request permission when needed, explain purpose, limit precision/retention, restrict access."), ("Administration", "Least privilege, strong authentication, audit trails, protected exports, periodic reviews."), ("Incident response", "Severity, containment, evidence, notification decision, recovery, and review.")], [2100, 7260], font_size=9.2)
    add_callout(doc, "Compliance note", "This specification identifies engineering controls, not legal advice. Obtain Sri Lankan legal review of privacy, consumer, electronic transaction, tax, food/product, payment, labor, and marketplace obligations.", fill="FFF6E6", accent=AMBER)
    add_heading(doc, "12. Deployment and Operations", 1)
    for step in ["Provision PHP 8.2+, MySQL 8, reverse proxy, private storage, queue manager, scheduler, TLS, and monitoring.", "Set production environment, APP_DEBUG=false, HTTPS URL, database, provider, Maps, PayHere, backup, and storage secrets.", "Install production dependencies and build Vite assets.", "Run reviewed migrations and create the storage link only for intended public assets.", "Optimize Laravel; start/restart workers; enable scheduler execution every minute.", "Verify health, auth, checkout, callbacks, notifications, map restrictions, jobs, backup, and restore.", "Promote mobile only after API contract, signing, privacy declarations, Firebase, and store checks pass."]: add_list_item(doc, step, number_id)
    add_heading(doc, "13. Acceptance and Verification", 1)
    add_table(doc, ["Area", "Acceptance evidence"], [("Identity/RBAC", "Tests and role matrix prove registration, verification, approval, token abilities, ownership, and denial."), ("Catalog/checkout", "Tests prove visibility, price consistency, vendor split, coupon limit, delivery fee, and oversell prevention."), ("Payments/refunds", "Sandbox callbacks, replay tests, wallet concurrency, COD, refund boundaries, reconciliation."), ("Delivery", "End-to-end state transition, proof storage, and failure tests."), ("Subscriptions", "Scheduler, retry, date calculation, stock failure, state, and generated order tests."), ("Mobile/API", "Every enabled screen passes staging contract tests; no missing route calls."), ("Security/privacy", "Threat model, audits, rate limits, secret/upload tests, data inventory, privacy review."), ("Operations", "Production-like deployment, load test, alerts, encrypted backup, timed restore drill.")], [1800, 7560], font_size=9.2)
    add_heading(doc, "14. Dependencies, Risks, and Open Decisions", 1)
    add_table(doc, ["ID", "Item", "Required decision or mitigation"], [("R-01", "Flutter calls many absent API endpoints and default URL omits /v1.", "Freeze API contract; implement or feature-gate; add contract tests."), ("R-02", "Mobile verification unavailable in current environment.", "Run flutter analyze/test and signed staging builds in CI."), ("R-03", "Some guides describe older statuses/schema.", "Treat migrations/services as authoritative; regenerate docs."), ("R-04", "Wallet top-up and payment simulations are placeholders.", "Disable or connect authorized production flows."), ("R-05", "Service targets are not formally agreed.", "Define load, availability, RPO/RTO, retention, support hours."), ("R-06", "Compliance incomplete without legal review.", "Complete privacy/consumer controls and review."), ("OD-01", "Supported launch geography", "Approve districts/zones, coverage, capacity, unavailable behavior."), ("OD-02", "Cancellation/refund policy", "Reconcile UI, service rules, obligations, partial refund, settlement."), ("OD-03", "Notification channels", "Approve providers, opt-in/out, templates, retries, cost, fallback.")], [1000, 3750, 4610], font_size=8.9)
    add_heading(doc, "Appendix A. Authoritative State Models", 1)
    add_table(doc, ["Object", "States / principal transitions"], [("Order", "pending -> confirmed -> packed -> assigned_to_rider -> out_for_delivery -> delivered; cancel/refund branches"), ("Delivery", "pending -> assigned -> picked_up -> on_the_way -> delivered; failed/cancelled branches"), ("Payment", "pending -> paid or failed; paid -> refunded; COD paid at delivery"), ("Vendor", "pending -> approved or rejected; suspension is administrative"), ("Rider", "verification pending -> verified/rejected/suspended; availability unavailable/available/delivering"), ("Product", "draft/pending -> approved; rejected/inactive/out_of_stock branches"), ("Subscription", "active <-> paused; -> cancelled; ended -> completed"), ("Support ticket", "open -> in_progress -> resolved -> closed")], [1800, 7560], font_size=9.2)
    add_heading(doc, "Appendix B. API Baseline", 1); add_para(doc, "The observed route table contains 298 routes: 27 API, 94 admin, 53 customer, 37 vendor, 15 rider, 19 super-admin, and 53 public/auth/shared. The API baseline is /api/v1.")
    add_table(doc, ["Scope", "Observed endpoints"], [("Public", "POST register/login; GET categories/products/product detail"), ("Common", "POST logout; GET profile; email/phone OTP send/verify"), ("Customer", "Cart CRUD; checkout quote; order create/list/show"), ("Rider", "Deliveries/detail; status update; location"), ("Vendor", "Overview, orders, wallet")], [1700, 7660], font_size=9.2)
    add_heading(doc, "Appendix C. Traceability Summary", 1)
    add_table(doc, ["Family", "Implementation evidence", "Verification"], [("FR-AUTH", "auth/API routes, middleware, provider, controllers", "Authentication, API security, registration tests"), ("FR-CAT/CHK", "catalog/cart/order services, controllers, requests, policies", "Core API and critical commerce tests"), ("FR-ORD/PAY", "order, delivery, payment, wallet, refund services", "Critical tests plus staging scenarios"), ("FR-ENG/SUB", "promotion, loyalty, notification, support, review, subscription", "Expanded feature and scheduler tests"), ("FR-ADM", "admin routes/controllers, reports, backup", "Backup test, role matrix, restore drill"), ("FR-API", "API routes/controllers/resources, Flutter services/providers", "API and proposed contract tests")], [1700, 4700, 2960], font_size=8.9)
    add_heading(doc, "Appendix D. Glossary", 1)
    add_table(doc, ["Term", "Definition"], [("COD", "Cash on Delivery."), ("LKR", "Sri Lankan Rupee currency code."), ("RBAC", "Role-based access control."), ("Sanctum", "Laravel token/session authentication package."), ("SKU", "Stock keeping unit."), ("RPO", "Recovery point objective."), ("RTO", "Recovery time objective."), ("Idempotent", "Safe to repeat without a second unintended effect."), ("Vendor split", "One persisted order per vendor from a multi-vendor cart."), ("Proof of delivery", "Photo, optional signature, note, and timestamp at completion.")], [1800, 7560], font_size=9.2)
    doc.save(SRS_PATH)


def build_research(architecture_path: Path, readiness_path: Path):
    doc = Document(); setup_styles(doc, "narrative"); configure_section(doc.sections[0], "DailyCart Technical Research and Project Analysis")
    doc.core_properties.title, doc.core_properties.subject, doc.core_properties.author = "DailyCart Technical Research and Project Analysis", "Technical research, feasibility, readiness, and roadmap", "DailyCart Project Team"
    bullet_id, number_id = add_numbering(doc, "bullet", text_left=540, hanging=270, after=80, line=300), add_numbering(doc, "decimal", text_left=540, hanging=270, after=80, line=300)
    add_cover_research(doc)
    add_heading(doc, "Abstract", 1)
    add_para(doc, "This study evaluates DailyCart as a multi-vendor grocery delivery system for Sri Lanka. The analysis combines repository inspection, route and schema inventory, source-level tracing of critical business rules, automated verification, web/mobile contract comparison, and comparison against authoritative guidance for Laravel, Flutter, MySQL, API security, PayHere, Google Maps, and Sri Lankan digital-commerce obligations. The web platform demonstrates unusually broad functional coverage: marketplace governance, stock-safe multi-vendor checkout, four payment modes, rider delivery, loyalty, subscriptions, refunds, support, reports, notifications, integrations, and encrypted backups. All 54 available automated tests passed with 241 assertions, and the web asset production build completed successfully.")
    add_para(doc, "The central finding is that the Laravel web application is a credible feature-complete baseline, but the end-to-end product is not yet release-complete. The Flutter client contains 227 Dart files and rich customer/vendor/rider screens, yet many client service paths do not exist in the current 27-route API surface, and the default mobile base URL omits the /v1 prefix. The recommended strategy is to preserve the modular monolith, establish an explicit OpenAPI contract, close the mobile/API gap, harden payment and top-up placeholders, complete privacy and consumer-transparency controls, define operational service levels, and add production-like load, recovery, and mobile verification.")
    add_callout(doc, "Research conclusion", "DailyCart is technically viable and architecturally appropriate for a controlled pilot. Production launch should be gated on API contract completion, payment hardening, privacy/compliance implementation, and operational validation.")
    add_heading(doc, "1. Research Context", 1); add_heading(doc, "1.1 Problem Setting", 2)
    add_para(doc, "Grocery delivery platforms coordinate short-lived stock, vendor ownership, geographic coverage, time-slot promises, last-mile handoff, payment uncertainty, refunds, consumer disclosures, and frequent status communication. A Sri Lankan platform also benefits from LKR-native pricing, local payment integration, district delivery logic, mobile-first access, and resilience to variable connectivity.")
    add_para(doc, "Sri Lanka's Department of Census and Statistics reported that 54.6 percent of the population aged 5-69 used the internet in 2024, and smartphones accounted for 80.3 percent of devices used to connect to internet or email [1]. A U.S. International Trade Administration guide reported significant local e-commerce growth and estimated that 43 percent of roughly 11 million internet users had made online purchases, while the Western Province accounted for around half of orders [2]. Central Bank reporting for 2025 recorded continued retail digital payment growth, with CEFTS representing 55 percent of retail payment value [3]. These indicators support demand but also make rollout geography, affordability, trust, and assisted recovery critical.")
    add_heading(doc, "1.2 Research Objectives", 2)
    for item in ["Determine whether the selected architecture fits current scale and complexity.", "Assess feature breadth, transaction integrity, security, testing, and operations.", "Compare the Laravel web/API contract with Flutter expectations.", "Identify technical, operational, economic, legal, privacy, and delivery risks.", "Provide a prioritized roadmap and measurable pilot/production criteria."]: add_list_item(doc, item, bullet_id)
    add_heading(doc, "1.3 Research Questions", 2)
    for item in ["RQ1. Does the modular monolith support marketplace, payment, delivery, and administration?", "RQ2. Are commerce invariants protected against concurrency, ownership, and partial failure?", "RQ3. Is the mobile contract aligned with the backend API?", "RQ4. Which gaps block pilot or production launch?", "RQ5. What evidence would credibly demonstrate readiness?"]: add_list_item(doc, item, number_id)
    add_heading(doc, "2. Methodology", 1)
    add_para(doc, "The study used a structured, evidence-led technical review. It did not infer completion from screen names or project guides alone. Where documentation and source disagreed, executable routes, migrations, services, policies, and tests were authoritative.")
    add_table(doc, ["Method", "Evidence", "Purpose"], [("Repository inventory", "51 models, 91 controllers, 35 services, 18 policies, 35 requests, 26 notifications, 189 Blade views, 227 Dart files", "Estimate scope and architecture"), ("Contract inventory", "298 Laravel routes, 27 API routes; Flutter service paths", "Compare server and client"), ("Rule tracing", "Cart, order, payment, wallet, refund, delivery, promotion, loyalty, subscription, notification, backup services", "Validate state invariants"), ("Data review", "Migrations, relationships, indexes, statuses, money, queues, logs, backups", "Assess integrity/operations"), ("Verification", "PHPUnit suite and Vite build", "Measure executable evidence"), ("External comparison", "Official technical, market, payment, mapping, security, and regulatory sources", "Benchmark decisions")], [1700, 4700, 2960], header_fill=NARRATIVE_FILL, font_size=8.9)
    add_para(doc, "Limitations: no production traffic, live PayHere merchant, live Maps credentials, signed mobile pipeline, penetration test, assistive-technology test, or legal opinion was available. The Flutter CLI could not be verified locally. Readiness scores are analytical, not certification.")
    add_heading(doc, "3. System Under Study", 1); add_figure(doc, architecture_path, "Figure 1. Current-state DailyCart architecture derived from the repository.")
    add_heading(doc, "3.1 Technology Profile", 2)
    add_table(doc, ["Layer", "Technology", "Observed role"], [("Backend", "PHP 8.2+, Laravel 12.64.0", "Routing, validation, policies, services, ORM, queues, scheduler, notifications"), ("Web", "Blade, Tailwind, Alpine, Axios, Vite", "Responsive public and role dashboards"), ("Data", "MySQL 8", "Transactional marketplace and operational persistence"), ("API auth", "Laravel Sanctum 4", "Expiring bearer tokens and abilities"), ("RBAC", "Spatie Permission 6", "Roles/permissions with middleware and policies"), ("Mobile", "Flutter/Dart, Riverpod, Dio, GoRouter", "Customer/vendor/rider screens and services"), ("Integrations", "PayHere, Google Maps, email, SMS, push/Firebase, S3", "Payment, location, communication, backup")], [1500, 2600, 5260], header_fill=NARRATIVE_FILL, font_size=9.0)
    add_heading(doc, "3.2 Functional Breadth", 2); add_para(doc, "The codebase covers the commerce loop rather than a catalog prototype. Five role areas connect to catalog governance, multi-vendor order splitting, payment, delivery, finance, refund, loyalty, subscriptions, reviews, support, content, analytics, backups, and mobile. This breadth is valuable but raises the cost of state consistency, documentation, contract testing, and release coordination.")
    add_heading(doc, "4. Findings", 1); add_figure(doc, readiness_path, "Figure 2. Repository-based readiness assessment; 5 represents strong evidence for reviewed scope.")
    add_heading(doc, "4.1 Architecture Fitness", 2)
    add_para(doc, "The modular monolith is a sound current choice. Core workflows share transactions across stock, orders, payments, deliveries, wallets, loyalty, refunds, and notifications. One Laravel deployment simplifies ACID behavior, authorization, development, and ownership. Premature microservices would add distributed transaction, ordering, observability, deployment, and data-ownership complexity without measured need.")
    add_para(doc, "The service layer is the strongest pattern. Laravel documents Sanctum as a preferred fit for first-party web plus mobile/token APIs [4], matching DailyCart's scoped abilities. Laravel also warns that queued work inside transactions should run after commit [5]; DailyCart uses after-commit queuing for several communications.")
    add_heading(doc, "4.2 Commerce Integrity", 2)
    add_para(doc, "Checkout and wallet behavior show mature consistency. Eligibility is rechecked, product rows are locked, order creation is transactional, and the cart converts only after all vendor orders exist. Wallet debit locks the customer, rejects overdraft, records balance-after, and marks payment paid. Refund approval coordinates wallet, payment, order, and loyalty in one transaction.")
    add_para(doc, "These choices match MySQL InnoDB locking reads [6]. Tests exercise competing checkout, quote/order parity, coupon limits, wallet single debit, and delivery fees. Remaining assurance should cover callback idempotency, recurring jobs, refund races, and repeated rider updates.")
    add_heading(doc, "4.3 Security Posture", 2)
    add_para(doc, "Positive controls include verification gates, token abilities, role middleware, ownership policies, rate limits, CSRF with a callback exception, upload validation, encrypted backups, dependency audits, and separate integration secrets. These address central OWASP API risks including object authorization, authentication, resource consumption, and sensitive flows [7].")
    add_para(doc, "Priority work includes threat modeling, negative authorization tests, callback replay protection, mobile authorization, disabling simulations/placeholders, secret scanning, security headers, cookie review, log redaction, upload malware strategy, access review, and incident response. Trusting all proxies should be constrained to known production proxies.")
    add_heading(doc, "4.4 Payment and Mapping Integrations", 2)
    add_para(doc, "PayHere states that notify_url is the authoritative server callback, localhost cannot receive it, and md5sig must be recomputed before state changes [8]. DailyCart's signature validation is directionally correct; production still needs idempotency, exact amount/currency/order matching, chargeback mapping, reconciliation, and sandbox/live separation.")
    add_para(doc, "Google recommends application/API key restrictions, separate keys per app/platform where practical, and server proxying for sensitive mobile web-service calls [9]. DailyCart separates browser, server, and mobile keys. Release should enforce restrictions, quotas, monitoring, HTTPS, and a plan to replace legacy Distance Matrix usage.")
    add_heading(doc, "4.5 Mobile/API Contract Gap", 2)
    add_para(doc, "This is the clearest blocker. Laravel exposes 27 API routes under /api/v1. Flutter defaults end at /api and services call /login or /products, so implemented endpoints resolve incorrectly without central v1 insertion. Flutter also expects addresses, wishlist, coupons, checkout, payment, notifications, loyalty, reviews, support, search, vendor products/profile/dashboard/earnings, and richer rider endpoints absent from routes/api.php.")
    add_table(doc, ["Area", "Backend", "Flutter", "Assessment"], [("Base URL", "/api/v1", "/api plus /login etc.", "Immediate mismatch"), ("Customer", "Cart, quote, create/list/show orders", "Also address, wishlist, coupon, checkout, cancel, payment, loyalty, review, support", "Partial"), ("Vendor", "Overview, orders, wallet", "Dashboard/profile/products/inventory/orders/earnings/reviews", "Large gap"), ("Rider", "Deliveries/detail/status/location", "Dashboard/profile/earnings/proof upload", "Partial gap"), ("Products", "List/detail/categories", "Featured/best/new/flash/recommended/search", "Gap or query consolidation"), ("Notifications", "No API routes", "List/read/read-all/delete/device token", "Missing")], [1450, 2500, 3370, 2040], header_fill=NARRATIVE_FILL, font_size=8.5)
    add_para(doc, "Flutter recommends clear UI/data layers with repositories/services and testable boundaries [10]. DailyCart has screens, providers, services, models, and routing, but an explicit repository interface would improve mocking, caching, offline behavior, contract migration, and error normalization.")
    add_heading(doc, "4.6 Verification and Operations", 2)
    add_para(doc, "The baseline is meaningful: 54 tests and 241 assertions passed, including registration, API security, cart isolation, variant ownership, pricing parity, configurable delivery fees, oversell prevention, coupon use, wallet debit, encrypted backup retention, offers, and location. Vite also built successfully, and CI covers PHPUnit, formatting, build, and audits.")
    add_para(doc, "Evidence is thinner for delivery states, callback permutations, refunds, subscriptions, support/review, reports, retries, mobile, accessibility, performance, and restores. Worker, scheduler, deployment, and backup instructions exist, but SLOs, RPO/RTO, alerts, capacity, and a restore record do not.")
    add_heading(doc, "4.7 Regulatory and Consumer Considerations", 2)
    add_para(doc, "Sri Lanka's Personal Data Protection Act applies to local processing and providers offering goods/services to people in Sri Lanka [11]. The Data Protection Authority announced operational dates beginning in March 2025 for major parts [12]. DailyCart therefore needs a data inventory, purpose analysis, notices, rights procedures, retention, safeguards, incident governance, and processor controls.")
    add_para(doc, "Consumer Affairs Authority Direction 91 covers e-commerce privacy, transparent data use, safeguards, prominent privacy policies, and appropriate cancellation opportunities [13]. DailyCart's policy, receipt, cancellation, support, and review functions are a base; legal review should reconcile disclosures, actual behavior, vendor duties, fees, refunds, safety, and complaints.")
    add_heading(doc, "5. Comparative Design Evaluation", 1)
    add_table(doc, ["Decision", "Selected", "Alternative", "Evaluation"], [("Architecture", "Laravel modular monolith", "Microservices", "Prefer selected approach while core workflows share transactions and one team operates the system."), ("Web", "Blade/Tailwind/Alpine", "React/Vue SPA", "Selected approach reduces API/deployment overhead; SPA only if interaction or team independence demands it."), ("Mobile", "Flutter", "Native / React Native", "Flutter fits reuse; backend contract is the constraint."), ("Database", "MySQL 8/InnoDB", "PostgreSQL / NoSQL", "MySQL supports relational commerce and locks; migration lacks a business case."), ("API auth", "Sanctum abilities", "OAuth2/Passport", "Sanctum is appropriate for first-party mobile; OAuth2 only for delegated third parties."), ("Queue", "Database queue", "Redis/SQS", "Suitable for pilot; migrate when throughput/isolation evidence justifies it.")], [1550, 2150, 2050, 3610], header_fill=NARRATIVE_FILL, font_size=8.5)
    add_heading(doc, "6. Feasibility Analysis", 1)
    add_table(doc, ["Dimension", "Assessment", "Rationale"], [("Technical", "High with conditions", "Core web architecture, rules, tests, and build are viable; mobile/API and assurance remain."), ("Operational", "Moderate", "Roles are clear; staffing, onboarding, exceptions, monitoring, and restores need procedures."), ("Economic", "Moderate to high", "Open-source stack limits licenses; costs are hosting, Maps, payments, messaging, storage, support, riders, stores."), ("Schedule", "Moderate", "Pilot follows focused hardening; full mobile breadth needs scope control."), ("Legal/privacy", "Moderate risk", "Obligations are identifiable; processes and formal review are needed."), ("Scalability", "Good for pilot/early growth", "Laravel/MySQL/queues scale; optimize from metrics before decomposition.")], [1500, 1700, 6160], header_fill=NARRATIVE_FILL, font_size=9.0)
    add_heading(doc, "7. SWOT Analysis", 1)
    add_table(doc, ["Strengths", "Weaknesses"], [("Broad scope; transactional rules; five-role RBAC; LKR/PayHere/location fit; tests; CI; encrypted backups.", "Mobile/API mismatch; limited mobile tests; documentation drift; placeholders; no SLO/RPO/RTO; incomplete privacy operations.")], [4680, 4680], header_fill=NARRATIVE_FILL, font_size=9.2)
    add_table(doc, ["Opportunities", "Threats"], [("Mobile adoption; vendor digitization; district expansion; subscriptions; loyalty; local payments; analytics; localization.", "Fraud; overselling; delivery failure; privacy breach; external API cost; store rejection; vendor quality; competition; regulation.")], [4680, 4680], header_fill=NARRATIVE_FILL, font_size=9.2)
    add_heading(doc, "8. Prioritized Gap Analysis", 1)
    add_table(doc, ["Priority", "Gap", "Impact", "Action"], [("P0", "Mobile URL/contract", "Core screens fail", "Publish OpenAPI; fix /api/v1; implement/gate paths; contract tests."), ("P0", "Payment/top-up boundaries", "Financial loss/state error", "Disable placeholders; idempotent PayHere reconciliation; authorized top-up."), ("P0", "Privacy/consumer operations", "Legal and trust risk", "Data map, notices, rights, retention, processors, legal review."), ("P1", "Critical workflow tests", "Regressions", "Delivery/refund/subscription/concurrency/callback tests."), ("P1", "Targets/observability", "Slow detection/recovery", "SLO/RPO/RTO, metrics, alerts, restore/failover runbooks."), ("P1", "Mobile CI/release", "Platform/store risk", "Analyze/test/build, signing, Firebase, staged rollout."), ("P2", "Documentation drift", "Developer confusion", "Generate route/OpenAPI/data documentation."), ("P2", "Accessibility/localization", "Reduced reach", "WCAG audit, localization, Sinhala/Tamil, low-bandwidth UX.")], [900, 2400, 2550, 3510], header_fill=NARRATIVE_FILL, font_size=8.5)
    add_heading(doc, "9. Recommended Delivery Roadmap", 1)
    add_table(doc, ["Phase", "Deliverable"], [("Weeks 1-2: Contract", "Inventory enabled screens/calls; OpenAPI; fix /api/v1; define pilot modules and schemas."), ("Weeks 3-5: API", "Implement/gate customer essentials, vendor/rider minimum, notifications, profiles, addresses, cancellation, proof."), ("Weeks 4-6: Finance", "Idempotent callback, amount/currency validation, reconciliation, chargebacks, production flags, top-up."), ("Weeks 5-7: Assurance", "Threat model, authorization matrix, workflow/concurrency tests, secret/upload audits."), ("Weeks 6-8: Compliance", "Notices, data map, retention, rights, vendor terms, fee/refund disclosures, legal review."), ("Weeks 8-10: Operations", "SLO/RPO/RTO, dashboards, alerts, load test, restore drill, incident/rollback runbooks."), ("Weeks 9-11: Mobile", "Flutter CI, device matrix, offline/error states, Firebase, signed staging builds, store declarations."), ("Week 12: Pilot", "Limited districts/vendors/riders, live support, reconciliation, metrics, flags, rollback criteria.")], [2600, 6760], header_fill=NARRATIVE_FILL, font_size=9.1)
    add_heading(doc, "10. Evaluation Framework", 1)
    add_table(doc, ["Outcome", "Measure", "Pilot target"], [("API", "Enabled mobile calls with passing contract tests", "100 percent"), ("Checkout", "Successful orders / valid attempts", ">=99 percent excluding abandonment"), ("Stock", "Negative or oversold stock", "0"), ("Payment", "Unreconciled callbacks or duplicate ledger effects", "0; daily reconciliation"), ("Delivery", "On-time and proof completeness", "District target; proof 100 percent"), ("Performance", "Server p95 catalog/checkout", "<=1.5 s / <=3 s under agreed load"), ("Reliability", "Availability and queue age", "Approved SLO; no unalerted backlog"), ("Recovery", "Restore drill", "Pass before pilot and quarterly"), ("Security", "Open critical/high findings", "0 critical; highs remediated/accepted"), ("Experience", "Completion, support, refund reason, crash-free users", "Baseline then sprint improvement")], [1900, 4300, 3160], header_fill=NARRATIVE_FILL, font_size=8.8)
    add_heading(doc, "11. Conclusion", 1)
    add_para(doc, "DailyCart has crossed from prototype to serious application platform. The repository demonstrates coherent Laravel architecture, broad role functionality, local market fit, and deliberate protection of high-risk commerce invariants. Passing tests and the production web build strengthen confidence in the web baseline.")
    add_para(doc, "The next phase should emphasize completion and evidence rather than features. Contract-first API work, financial hardening, mobile verification, privacy/consumer implementation, and operational measurement can convert the existing breadth into a controlled pilot. The modular monolith remains the best balance of safety, velocity, and simplicity until measured scaling or organizational boundaries justify decomposition.")
    add_heading(doc, "References", 1)
    refs = [("[1]", "Sri Lanka Department of Census and Statistics. Labour Force Survey Annual Report 2024.", "https://www.statistics.gov.lk/Resource/en/LabourForce/Annual_Reports/LFS2024.pdf"), ("[2]", "U.S. International Trade Administration. Sri Lanka - eCommerce.", "https://www.trade.gov/country-commercial-guides/sri-lanka-ecommerce"), ("[3]", "Central Bank of Sri Lanka. Annual Economic Review 2025.", "https://www.cbsl.gov.lk/sites/default/files/cbslweb_documents/publications/aer/2025/en/Full_Text.pdf"), ("[4]", "Laravel. Authentication, version 12.x.", "https://laravel.com/docs/12.x/authentication"), ("[5]", "Laravel. Queues, version 12.x.", "https://laravel.com/docs/12.x/queues"), ("[6]", "Oracle MySQL. InnoDB transaction isolation and locking reads.", "https://dev.mysql.com/doc/refman/8.0/en/innodb-transaction-isolation-levels.html"), ("[7]", "OWASP. API Security Top 10 - 2023.", "https://owasp.org/API-Security/editions/2023/en/0x00-header/"), ("[8]", "PayHere. Checkout API and callback signature verification.", "https://support.payhere.lk/api-%26-mobile-sdk/checkout-api"), ("[9]", "Google. Maps Platform API security best practices.", "https://developers.google.com/maps/api-security-best-practices"), ("[10]", "Flutter. Guide to app architecture.", "https://docs.flutter.dev/app-architecture/guide"), ("[11]", "Government of Sri Lanka. Personal Data Protection Act, No. 9 of 2022.", "https://www.documents.gov.lk/view/act/2022/3/09-2022_E.pdf"), ("[12]", "Data Protection Authority of Sri Lanka. PDPA operational dates.", "https://www.dpa.gov.lk/est.php"), ("[13]", "Consumer Affairs Authority. Direction 91 for e-commerce.", "https://caa.gov.lk/web/images/Direction/English/sec10/91_E.pdf")]
    add_table(doc, ["Ref.", "Source", "URL"], refs, [800, 5000, 3560], header_fill=NARRATIVE_FILL, font_size=8.2)
    add_heading(doc, "Appendix A. Repository Evidence Summary", 1)
    add_table(doc, ["Evidence", "Observed result"], [("Routes", "298 total: 94 admin, 53 customer, 53 public/auth/shared, 37 vendor, 27 API, 19 super-admin, 15 rider"), ("Backend", "51 models, 91 controllers, 35 services, 18 policies, 35 requests, 26 notifications"), ("Presentation", "189 Blade views; Vite production build succeeded with 56 modules"), ("Mobile", "227 Dart files; role screens/providers/services/models/widgets; release verification unavailable"), ("Tests", "54 passed, 241 assertions, about 166 seconds"), ("Operations", "Database queues, subscription scheduler, token pruning, /up, encrypted backups, CI audits")], [2100, 7260], header_fill=NARRATIVE_FILL, font_size=9.0)
    doc.save(RESEARCH_PATH)


def audit_doc(path: Path):
    doc = Document(path)
    for section in doc.sections:
        assert round(section.page_width.inches, 2) == 8.5 and round(section.page_height.inches, 2) == 11
        assert round(section.left_margin.inches, 2) == 1 and round(section.right_margin.inches, 2) == 1
    for table in doc.tables:
        tbl_pr = table._tbl.tblPr; tbl_w = tbl_pr.find(qn("w:tblW")); layout = tbl_pr.find(qn("w:tblLayout"))
        assert tbl_w is not None and tbl_w.get(qn("w:type")) == "dxa" and layout is not None and layout.get(qn("w:type")) == "fixed"
        widths = [int(c.get(qn("w:w"))) for c in table._tbl.tblGrid.findall(qn("w:gridCol"))]; assert sum(widths) == 9360
        for row in table.rows:
            for idx, cell in enumerate(row.cells):
                tc_w = cell._tc.get_or_add_tcPr().find(qn("w:tcW")); assert tc_w is not None and int(tc_w.get(qn("w:w"))) == widths[idx]
    return {"paragraphs": len(doc.paragraphs), "tables": len(doc.tables), "sections": len(doc.sections)}


def main():
    OUT.mkdir(parents=True, exist_ok=True); QA.mkdir(parents=True, exist_ok=True)
    architecture, readiness = QA / "dailycart_architecture.png", QA / "dailycart_readiness.png"
    create_architecture_figure(architecture); create_readiness_figure(readiness)
    build_srs(architecture); build_research(architecture, readiness)
    print(SRS_PATH); print(audit_doc(SRS_PATH)); print(RESEARCH_PATH); print(audit_doc(RESEARCH_PATH))


if __name__ == "__main__": main()
